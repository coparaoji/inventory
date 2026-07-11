"""
Word template inspection and shape-fill injection.

Works at the raw OOXML (Open Packaging) level rather than through python-docx so
it can handle pre-drawn shape "slots" (``<wps:wsp>``) that python-docx does not
expose. Two capabilities:

* :func:`analyse_docx` / :func:`extract_media` -- inspect an existing .docx:
  list its zip contents, relationships, media, tables, drawings, shapes and
  image placements.
* :func:`fill_shape` -- fill an empty shape slot in a template with an image by
  swapping ``<a:solidFill>`` for ``<a:blipFill>`` and wiring up the media file
  and relationship.
"""
from __future__ import annotations

import copy
import zipfile
from pathlib import Path

import docx
from lxml import etree

# ---------------------------------------------------------------------------
# Namespaces
# ---------------------------------------------------------------------------

NS = {
    "w":    "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r":    "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rels": "http://schemas.openxmlformats.org/package/2006/relationships",
    "wp":   "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a":    "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic":  "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "wps":  "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "v":    "urn:schemas-microsoft-com:vml",
    "o":    "urn:schemas-microsoft-com:office:office",
}
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

_A = NS["a"]
_R = R_NS
_WPS = NS["wps"]
_RELS_NS = NS["rels"]
_RELS_TYPE_IMAGE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
)
_A14 = "http://schemas.microsoft.com/office/drawing/2010/main"


def _emu_to_in(emu: int) -> float:
    return round(emu / 914400, 3)


def _section(title: str, width: int = 56) -> None:
    print(f"\n  {'-' * width}")
    print(f"  {title}")
    print(f"  {'-' * width}")


# ---------------------------------------------------------------------------
# Inspection
# ---------------------------------------------------------------------------

def analyse_docx(path: Path) -> None:
    """Print a full inspection report for a single .docx file.

    Reports zip contents, relationships, media files, tables, drawings,
    shapes, VML shapes and image placements. Useful for reverse-engineering
    how a Word template is structured before filling it.
    """
    path = Path(path)
    print(f"\n{'=' * 70}")
    print(f"  FILE: {path.name}   ({path.stat().st_size / 1024:.1f} KB)")
    print(f"{'=' * 70}")

    with zipfile.ZipFile(path) as z:
        all_names = z.namelist()
        doc_bytes = z.read("word/document.xml")
        rels_bytes = z.read("word/_rels/document.xml.rels")

    doc_xml = etree.fromstring(doc_bytes)
    rels_xml = etree.fromstring(rels_bytes)

    _section("ZIP CONTENTS")
    with zipfile.ZipFile(path) as z:
        for info in sorted(z.infolist(), key=lambda i: i.filename):
            print(f"    {info.filename:<52}  {info.file_size / 1024:6.1f} KB")

    _section("RELATIONSHIPS")
    for rel in rels_xml.xpath("//rels:Relationship", namespaces=NS):
        rid = rel.get("Id")
        rtype = rel.get("Type", "").split("/")[-1]
        target = rel.get("Target")
        print(f"    {rid:<8}  {rtype:<22}  {target}")

    _section("MEDIA FILES")
    media = [n for n in all_names if n.startswith("word/media/")]
    if media:
        with zipfile.ZipFile(path) as z:
            for name in sorted(media):
                kb = z.getinfo(name).file_size / 1024
                print(f"    {name:<50}  {kb:7.1f} KB")
    else:
        print("    (none)")

    _section("TABLES")
    doc_obj = docx.Document(str(path))
    if not doc_obj.tables:
        print("    (none)")
    for t_idx, tbl in enumerate(doc_obj.tables):
        style = tbl.style.name if tbl.style else "?"
        print(f"    Table {t_idx}  -  {len(tbl.rows)}r x {len(tbl.columns)}c  style={style!r}")
        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                n_img = len(cell._tc.xpath(".//a:blip", namespaces=NS))
                text = cell.text.strip()[:40]
                if n_img or text:
                    print(f"      [{r_idx},{c_idx}]  images={n_img}  text={text!r}")

    _section("DRAWINGS  (<w:drawing>)")
    drawings = doc_xml.xpath("//w:drawing", namespaces=NS)
    print(f"    Total: {len(drawings)}")
    for i, drw in enumerate(drawings):
        doc_pr = drw.xpath(".//wp:docPr", namespaces=NS)
        name = doc_pr[0].get("name", "?") if doc_pr else "?"
        elem_id = doc_pr[0].get("id", "?") if doc_pr else "?"
        extent = drw.xpath(".//wp:extent", namespaces=NS)
        size_str = (
            f"{_emu_to_in(int(extent[0].get('cx', 0)))}\" x "
            f"{_emu_to_in(int(extent[0].get('cy', 0)))}\""
            if extent else "?"
        )
        kind = ("inline" if drw.xpath(".//wp:inline", namespaces=NS) else
                "anchored" if drw.xpath(".//wp:anchor", namespaces=NS) else "?")
        blips = drw.xpath(".//a:blip", namespaces=NS)
        embed = blips[0].get(f"{{{R_NS}}}embed", "?") if blips else "-"
        is_shp = bool(drw.xpath(".//wps:wsp", namespaces=NS))
        print(f"    [{i:>3}]  id={elem_id:<12} name={name!r:<32} "
              f"kind={kind:<9} size={size_str:<18} rId={embed:<8} shape={is_shp}")

    _section("SHAPES  (<wps:wsp>)")
    shapes = doc_xml.xpath("//wps:wsp", namespaces=NS)
    print(f"    Total: {len(shapes)}")
    for i, sp in enumerate(shapes):
        cnv2 = sp.xpath("ancestor::w:drawing//wp:docPr", namespaces=NS)
        name = cnv2[0].get("name", "?") if cnv2 else "?"
        prst_geom = sp.xpath(".//a:prstGeom", namespaces=NS)
        geom = prst_geom[0].get("prst", "?") if prst_geom else "custom"
        blip_f = sp.xpath(".//a:blipFill", namespaces=NS)
        solid_f = sp.xpath(".//a:solidFill", namespaces=NS)
        no_f = sp.xpath(".//a:noFill", namespaces=NS)
        if blip_f:
            blip = blip_f[0].xpath(".//a:blip", namespaces=NS)
            rid = blip[0].get(f"{{{R_NS}}}embed", "?") if blip else "?"
            fill = f"picture(rId={rid})"
        elif solid_f:
            fill = "solid"
        elif no_f:
            fill = "no-fill"
        else:
            fill = "?"
        xfrm = sp.xpath(".//a:xfrm/a:ext", namespaces=NS)
        size_str = (
            f"{_emu_to_in(int(xfrm[0].get('cx', 0)))}\" x "
            f"{_emu_to_in(int(xfrm[0].get('cy', 0)))}\""
            if xfrm else "?"
        )
        print(f"    [{i:>3}]  name={name!r:<32}  geom={geom:<12}  "
              f"fill={fill:<22}  size={size_str}")

    _section("VML SHAPES  (<v:shape>)")
    vml = doc_xml.xpath("//v:shape", namespaces=NS)
    if not vml:
        print("    (none)")
    for i, sp in enumerate(vml):
        name = sp.get("id", "?")
        style = sp.get("style", "")[:55]
        img_data = sp.xpath(".//v:imagedata", namespaces=NS)
        fill_info = (f"image rId={img_data[0].get(f'{{{R_NS}}}id', '?')}"
                     if img_data else f"filled={sp.get('filled', '?')}")
        print(f"    [{i:>2}]  id={name!r:<25}  {fill_info}  style={style!r}")

    _section("IMAGE PLACEMENTS  (a:blip -> rId -> file)")
    rid_map = {r.get("Id"): r.get("Target")
               for r in rels_xml.xpath("//rels:Relationship", namespaces=NS)}
    blips = doc_xml.xpath("//a:blip", namespaces=NS)
    if not blips:
        print("    (none)")
    for i, blip in enumerate(blips):
        rid = blip.get(f"{{{R_NS}}}embed") or blip.get(f"{{{R_NS}}}link", "?")
        fname = rid_map.get(rid, "???")
        doc_pr = None
        for anc in blip.iterancestors():
            if anc.tag in (f"{{{NS['wp']}}}inline", f"{{{NS['wp']}}}anchor"):
                kids = list(anc.iterchildren(f"{{{NS['wp']}}}docPr"))
                if kids:
                    doc_pr = kids[0]
                break
        shape_name = doc_pr.get("name", "?") if doc_pr else "?"
        print(f"    [{i:>3}]  rId={rid:<8}  file={fname:<35}  shape={shape_name!r}")


def extract_media(path: Path, media_root: Path) -> list[Path]:
    """Copy every ``word/media/*`` file from *path* into *media_root*.

    Files land under ``<media_root>/<docx-stem>/<zip-internal-path>`` so the
    original package layout is preserved. Returns the list of written paths.
    """
    path = Path(path)
    media_root = Path(media_root)
    written: list[Path] = []
    with zipfile.ZipFile(path) as z:
        media_entries = [n for n in z.namelist() if n.startswith("word/media/")]
        if not media_entries:
            return written
        doc_dir = media_root / path.stem
        for zip_path in sorted(media_entries):
            dest = doc_dir / Path(zip_path)
            dest.parent.mkdir(parents=True, exist_ok=True)
            dest.write_bytes(z.read(zip_path))
            written.append(dest)
    return written


# ---------------------------------------------------------------------------
# Shape fill injection
# ---------------------------------------------------------------------------
#
# The difference between an empty slot (solidFill) and a filled slot (blipFill)
# is one element swap inside <wps:spPr>. Three things must happen together:
#   1. DOCUMENT XML  - swap <a:solidFill> -> <a:blipFill r:embed="rIdN">
#   2. MEDIA FILE    - copy the image into word/media/ inside the zip
#   3. RELATIONSHIPS - add <Relationship Id="rIdN" Type=".../image" .../>
# The <w:drawing> parent (position, size, docPr name) is left untouched.


def _next_rid(rels_xml) -> str:
    """Return the next unused rId string (rId1, rId2, ...)."""
    existing = {
        r.get("Id")
        for r in rels_xml.xpath("//rels:Relationship",
                                namespaces={"rels": _RELS_NS})
    }
    n = 1
    while f"rId{n}" in existing:
        n += 1
    return f"rId{n}"


def _build_blip_fill(rid: str) -> "etree._Element":
    """Construct the <a:blipFill> element that replaces <a:solidFill>."""
    bf = etree.Element(f"{{{_A}}}blipFill", dpi="0", rotWithShape="1")
    blip = etree.SubElement(
        bf, f"{{{_A}}}blip", {f"{{{_R}}}embed": rid}, cstate="print"
    )
    # Word adds a useLocalDpi extension; include it for full compatibility.
    ext_lst = etree.SubElement(blip, f"{{{_A}}}extLst")
    ext = etree.SubElement(
        ext_lst, f"{{{_A}}}ext", uri="{28A0092B-C50C-407E-A947-70E740481C1C}"
    )
    etree.SubElement(ext, f"{{{_A14}}}useLocalDpi", {f"{{{_A14}}}val": "0"})

    etree.SubElement(bf, f"{{{_A}}}srcRect")
    stretch = etree.SubElement(bf, f"{{{_A}}}stretch")
    etree.SubElement(stretch, f"{{{_A}}}fillRect")
    return bf


def fill_shape(template: Path, output: Path, shape_index: int, image: Path) -> None:
    """Fill the shape at *shape_index* in *template* with *image*.

    Reads *template* without modifying it and writes the result to *output*.
    Raises :class:`IndexError` if *shape_index* is out of range.
    """
    template, output, image = Path(template), Path(output), Path(image)

    with zipfile.ZipFile(template) as z:
        files = {name: z.read(name) for name in z.namelist()}

    doc_root = etree.fromstring(files["word/document.xml"])
    rels_root = etree.fromstring(files["word/_rels/document.xml.rels"])

    rid = _next_rid(rels_root)
    media_zip_path = f"word/media/{image.name}"
    files[media_zip_path] = image.read_bytes()

    etree.SubElement(
        rels_root,
        "Relationship",
        Id=rid,
        Type=_RELS_TYPE_IMAGE,
        Target=f"media/{image.name}",
    )

    shapes = doc_root.xpath("//wps:wsp", namespaces={"wps": _WPS})
    if shape_index >= len(shapes):
        raise IndexError(
            f"shape_index {shape_index} out of range "
            f"(document has {len(shapes)} shapes)"
        )
    sp = shapes[shape_index]
    sp_pr = sp.find(f"{{{_WPS}}}spPr")
    solid = sp_pr.find(f"{{{_A}}}solidFill")
    if solid is not None:
        idx_in_parent = list(sp_pr).index(solid)
        sp_pr.remove(solid)
        sp_pr.insert(idx_in_parent, _build_blip_fill(rid))
    else:
        # Already has a blipFill -- just update the embed reference.
        blip = sp_pr.find(f".//{{{_A}}}blip")
        if blip is not None:
            blip.set(f"{{{_R}}}embed", rid)

    files["word/document.xml"] = etree.tostring(
        doc_root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    files["word/_rels/document.xml.rels"] = etree.tostring(
        rels_root, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
