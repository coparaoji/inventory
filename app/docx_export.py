"""
DOCX grid export.

Each product type gets its own page section (with the correct orientation)
and a table grid of images sized to the product's physical print dimensions.
Images are resized via Pillow and saved as PNG before insertion.
"""
from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path

from PIL import Image as PilImage, ImageDraw
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Product type registry
# ---------------------------------------------------------------------------

LIGHTER_COLOR_TYPES: set[str] = {"Lighter (White)", "Lighter (Gold)", "Lighter (Silver)"}
LIGHTER_TYPES: set[str]       = {"Lighter"} | LIGHTER_COLOR_TYPES

PRODUCT_TYPES = [
    "Lighter",
    "Lighter (White)",
    "Lighter (Gold)",
    "Lighter (Silver)",
    "Stash Jar",
    "Grinder",
    "Rolling Tray",
    "Other",
]

LIGHTER_COLOR_LABELS: dict[str, str] = {
    "Lighter (White)":  "White",
    "Lighter (Gold)":   "Gold",
    "Lighter (Silver)": "Silver",
}


@dataclass
class ProductSpec:
    """Target print dimensions and layout for one product category."""
    width_in: float
    height_in: float
    cols: int
    landscape: bool = False
    corner_radius: int = 0
    circle: bool = False
    border_px: int = 0
    border_color: tuple = (0, 0, 0, 255)


_LIGHTER_BASE = dict(width_in=1.42, height_in=2.22, cols=5, corner_radius=15)

SPECS: dict[str, ProductSpec] = {
    "Lighter":          ProductSpec(**_LIGHTER_BASE),
    "Lighter (White)":  ProductSpec(**_LIGHTER_BASE, border_px=8, border_color=(210, 210, 210, 255)),
    "Lighter (Gold)":   ProductSpec(**_LIGHTER_BASE, border_px=8, border_color=(212, 175, 55, 255)),
    "Lighter (Silver)": ProductSpec(**_LIGHTER_BASE, border_px=8, border_color=(192, 192, 192, 255)),
    "Stash Jar":        ProductSpec(width_in=2.55, height_in=2.55, cols=3, circle=True),
    "Grinder":          ProductSpec(width_in=2.35, height_in=2.35, cols=3),
    "Rolling Tray":     ProductSpec(width_in=4.00, height_in=5.70, cols=2),
    "Other":            ProductSpec(width_in=2.00, height_in=2.00, cols=3),
}

PAGE_W  = 8.5    # portrait page width  (inches)
PAGE_H  = 11.0   # portrait page height (inches)
MARGIN  = 0.0    # all-sides margin     (inches)

# Rolling Tray exact print dimensions
TRAY_LANDSCAPE_W = 5.70   # landscape slot: wide side
TRAY_LANDSCAPE_H = 4.00   # landscape slot: short side
TRAY_PORTRAIT_W  = 4.00   # each portrait slot: short side
TRAY_PORTRAIT_H  = 5.70   # each portrait slot: tall side


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

@dataclass
class QueueEntry:
    path: Path
    product_type: str = "Other"
    quantity: int = 1


def build_docx(entries: list[QueueEntry], dest: Path) -> None:
    """Build the multi-section grid DOCX and save it to *dest*."""
    doc = Document()

    groups: dict[str, list[QueueEntry]] = {}
    for e in entries:
        groups.setdefault(e.product_type, []).append(e)

    for idx, (product_type, group) in enumerate(groups.items()):
        spec = SPECS.get(product_type, SPECS["Other"])

        section = doc.sections[0] if idx == 0 else doc.add_section()
        _configure_section(section, spec.landscape)

        # Expand quantity into repeated paths; lighters get 2 images per unit (two sides)
        image_paths: list[Path] = []
        for e in group:
            qty = max(1, e.quantity)
            if product_type in LIGHTER_TYPES:
                image_paths.extend([e.path, e.path] * qty)
            else:
                image_paths.extend([e.path] * qty)

        if product_type == "Rolling Tray":
            _add_tray_layout(doc, image_paths)
        else:
            actual_w = _fit_width(spec)
            _add_image_grid(doc, image_paths, spec, actual_w)
            # Color label placed as plain text right after the grid (no footer)
            if product_type in LIGHTER_COLOR_TYPES:
                label = LIGHTER_COLOR_LABELS[product_type]
                p = doc.add_paragraph(label)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _zero_para_spacing(p)
                run = p.runs[0]
                run.font.size = Pt(14)
                run.font.bold = True

    doc.save(str(dest))


# ---------------------------------------------------------------------------
# Internals
# ---------------------------------------------------------------------------

def _fit_width(spec: ProductSpec) -> float:
    """Largest image width (inches) that fits within the page's column layout."""
    page_content_w = (PAGE_H if spec.landscape else PAGE_W) - 2 * MARGIN
    per_col = page_content_w / spec.cols
    return min(spec.width_in, per_col * 0.96)


def _configure_section(section, landscape: bool) -> None:
    m = Inches(MARGIN)
    section.top_margin    = m
    section.bottom_margin = m
    section.left_margin   = m
    section.right_margin  = m

    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width  = Inches(PAGE_H)
        section.page_height = Inches(PAGE_W)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width  = Inches(PAGE_W)
        section.page_height = Inches(PAGE_H)


def _zero_cell_margins(cell) -> None:
    """Set all internal cell padding to zero."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), "0")
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _zero_para_spacing(para) -> None:
    """Remove paragraph space-before / space-after."""
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)


def _set_row_height(row, height_in: float) -> None:
    """Force an exact row height (in inches) via OOXML."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_in * 1440)))  # 1440 twips per inch
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


# ---------------------------------------------------------------------------
# Standard image grid (used by all product types except Rolling Tray)
# ---------------------------------------------------------------------------

def _add_image_grid(
    doc: Document,
    paths: list[Path],
    spec: ProductSpec,
    actual_w: float,
) -> None:
    if not paths:
        return

    cols = spec.cols
    rows_needed = (len(paths) + cols - 1) // cols

    table = doc.add_table(rows=rows_needed, cols=cols)
    table.style = "Table Grid"

    for i, img_path in enumerate(paths):
        row_idx = i // cols
        col_idx = i % cols
        cell = table.cell(row_idx, col_idx)
        cell.width = Inches(actual_w)
        _zero_cell_margins(cell)

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _zero_para_spacing(para)
        run = para.add_run()

        try:
            buf = _prepare_image(img_path, actual_w, spec.height_in, spec=spec)
            run.add_picture(buf, width=Inches(actual_w), height=Inches(spec.height_in))
        except Exception:
            run.text = f"[Error: {img_path.name}]"

    # Clear leftover empty cells in the last row
    for empty in range(len(paths), rows_needed * cols):
        cell = table.cell(empty // cols, empty % cols)
        _zero_cell_margins(cell)
        for para in cell.paragraphs:
            para.clear()
            _zero_para_spacing(para)


# ---------------------------------------------------------------------------
# Rolling Tray layout  (portrait page, 1 landscape top + 2 portrait bottom)
# ---------------------------------------------------------------------------

def _add_tray_layout(doc: Document, paths: list[Path]) -> None:
    """Consume *paths* in chunks of 3 — one page per chunk.
    Slot order: index 0 → landscape top, index 1 → portrait bottom-left,
    index 2 → portrait bottom-right.  Incomplete last chunks leave slots empty.
    """
    if not paths:
        return
    chunks = [paths[i:i + 3] for i in range(0, len(paths), 3)]
    for page_idx, chunk in enumerate(chunks):
        if page_idx > 0:
            _add_page_break(doc)
        _add_tray_page(doc, chunk)


def _add_tray_page(doc: Document, paths: list[Path]) -> None:
    """Render one tray page from up to 3 paths (landscape, portrait-L, portrait-R)."""
    landscape_path  = paths[0] if len(paths) > 0 else None
    portrait_left   = paths[1] if len(paths) > 1 else None
    portrait_right  = paths[2] if len(paths) > 2 else None

    # ── Landscape image centred on its own paragraph ────────────────────────
    top_para = doc.add_paragraph()
    top_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _zero_para_spacing(top_para)
    if landscape_path:
        try:
            buf = _prepare_image(landscape_path, TRAY_LANDSCAPE_W, TRAY_LANDSCAPE_H)
            top_para.add_run().add_picture(
                buf, width=Inches(TRAY_LANDSCAPE_W), height=Inches(TRAY_LANDSCAPE_H)
            )
        except Exception:
            top_para.add_run().text = "[Error]"

    # ── Two portrait images side by side ────────────────────────────────────
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(TRAY_PORTRAIT_W)
    table.columns[1].width = Inches(TRAY_PORTRAIT_W)
    _set_row_height(table.rows[0], TRAY_PORTRAIT_H)

    for col, port_path in enumerate([portrait_left, portrait_right]):
        cell = table.cell(0, col)
        cell.width = Inches(TRAY_PORTRAIT_W)
        _zero_cell_margins(cell)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _zero_para_spacing(para)
        if port_path:
            try:
                buf = _prepare_image(port_path, TRAY_PORTRAIT_W, TRAY_PORTRAIT_H)
                para.add_run().add_picture(
                    buf, width=Inches(TRAY_PORTRAIT_W), height=Inches(TRAY_PORTRAIT_H)
                )
            except Exception:
                para.add_run().text = "[Error]"


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    _zero_para_spacing(p)
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ---------------------------------------------------------------------------
# Image processing
# ---------------------------------------------------------------------------

def _apply_rounded_corners(img: PilImage.Image, radius: int = 20) -> PilImage.Image:
    img = img.convert("RGBA")
    mask = PilImage.new("L", img.size, 0)
    ImageDraw.Draw(mask).rounded_rectangle([(0, 0), img.size], radius=radius, fill=255)
    img.putalpha(mask)
    return img


def _apply_border(
    img: PilImage.Image,
    border_px: int = 6,
    radius: int = 20,
    color: tuple = (0, 0, 0, 255),
) -> PilImage.Image:
    img = img.convert("RGBA")
    w, h = img.size
    half = border_px // 2
    ImageDraw.Draw(img).rounded_rectangle(
        [(half, half), (w - half, h - half)],
        radius=radius,
        outline=color,
        width=border_px,
    )
    return img


def _prepare_image(
    path: Path,
    width_in: float,
    height_in: float,
    dpi: int = 150,
    spec: ProductSpec | None = None,
) -> io.BytesIO:
    """
    Resize *path* to exact target dimensions and return as a PNG BytesIO.
    Orientation is auto-aligned: the longer source side maps to the longer
    target side. Aspect-ratio distortion is intentional for print fit.
    """
    target_w = max(1, int(width_in  * dpi))
    target_h = max(1, int(height_in * dpi))

    border_px    = spec.border_px    if spec else 0
    border_color = spec.border_color if spec else (0, 0, 0, 255)

    if spec and spec.circle:
        corner_radius = min(target_w, target_h) // 2
    else:
        corner_radius = spec.corner_radius if spec else 0

    with PilImage.open(path) as img:
        img = img.convert("RGBA")
        src_w, src_h = img.size
        # Align orientation: if source and target disagree on which side is longer,
        # rotate the source 90° so the longer axis matches — output stays (target_w, target_h).
        if (src_w > src_h) != (target_w > target_h):
            img = img.rotate(90, expand=True)
        img = img.resize((target_w, target_h), PilImage.Resampling.LANCZOS)
        already_transparent = img.split()[3].getextrema()[0] < 255
        if corner_radius > 0 and not already_transparent:
            img = _apply_rounded_corners(img, radius=corner_radius)
        if border_px > 0:
            img = _apply_border(img, border_px=border_px, radius=corner_radius, color=border_color)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf
