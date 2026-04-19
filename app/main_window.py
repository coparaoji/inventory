import io
import os
from pathlib import Path

from PIL import Image as PilImage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PyQt6.QtWidgets import (
    QMainWindow,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QPushButton,
    QTableWidget,
    QTableWidgetItem,
    QLineEdit,
    QLabel,
    QStatusBar,
    QHeaderView,
    QFileDialog,
    QListWidget,
    QListWidgetItem,
    QFrame,
    QMessageBox,
    QAbstractItemView,
    QSplitter,
)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QColor

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tiff", ".tif"}

GRID_COLUMNS = 3
# Max image size (pixels) inside each docx cell before inserting
THUMB_PX = 200
# Column width in inches for the docx table
COL_WIDTH_IN = 2.2


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Inventory")
        self.setMinimumSize(980, 680)
        self._current_folder: Path | None = None
        self._queue: list[Path] = []
        self._build_ui()

    # ------------------------------------------------------------------
    # UI construction
    # ------------------------------------------------------------------

    def _build_ui(self):
        central = QWidget()
        self.setCentralWidget(central)

        root = QVBoxLayout(central)
        root.setContentsMargins(12, 12, 12, 12)
        root.setSpacing(8)

        splitter = QSplitter(Qt.Orientation.Vertical)
        root.addWidget(splitter, stretch=1)

        # ---- top half: browser ----
        browser_widget = QWidget()
        browser_layout = QVBoxLayout(browser_widget)
        browser_layout.setContentsMargins(0, 0, 0, 0)
        browser_layout.setSpacing(6)

        # folder row
        folder_row = QHBoxLayout()
        self.folder_label = QLabel("No folder selected")
        self.folder_label.setStyleSheet("color: grey;")
        folder_btn = QPushButton("Select Folder…")
        folder_btn.clicked.connect(self._on_select_folder)
        folder_row.addWidget(folder_btn)
        folder_row.addWidget(self.folder_label, stretch=1)
        browser_layout.addLayout(folder_row)

        # search row
        search_row = QHBoxLayout()
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Filter by filename…")
        self.search_input.textChanged.connect(self._on_search)
        search_row.addWidget(QLabel("Search:"))
        search_row.addWidget(self.search_input, stretch=1)

        add_btn = QPushButton("Add Selected to Queue →")
        add_btn.setToolTip("Add highlighted images to the print queue")
        add_btn.clicked.connect(self._on_add_to_queue)
        search_row.addWidget(add_btn)
        browser_layout.addLayout(search_row)

        # file table
        self.table = QTableWidget(0, 3)
        self.table.setHorizontalHeaderLabels(["Filename", "Extension", "Size"])
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        self.table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        self.table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.table.setSortingEnabled(True)
        browser_layout.addWidget(self.table, stretch=1)

        splitter.addWidget(browser_widget)

        # ---- bottom half: print queue ----
        queue_widget = QWidget()
        queue_layout = QVBoxLayout(queue_widget)
        queue_layout.setContentsMargins(0, 4, 0, 0)
        queue_layout.setSpacing(6)

        # divider label + controls
        queue_header = QHBoxLayout()
        queue_title = QLabel("Print Queue")
        queue_title.setStyleSheet("font-weight: bold; font-size: 13px;")
        self.queue_count_label = QLabel("0 image(s)")
        self.queue_count_label.setStyleSheet("color: grey;")

        clear_btn = QPushButton("Clear Queue")
        clear_btn.clicked.connect(self._on_clear_queue)

        remove_btn = QPushButton("Remove Selected")
        remove_btn.clicked.connect(self._on_remove_from_queue)

        self.print_btn = QPushButton("Print to DOCX…")
        self.print_btn.setEnabled(False)
        self.print_btn.setStyleSheet(
            "QPushButton { background-color: #2d7d46; color: white; font-weight: bold; padding: 4px 12px; }"
            "QPushButton:disabled { background-color: #555; color: #999; }"
        )
        self.print_btn.clicked.connect(self._on_print_docx)

        queue_header.addWidget(queue_title)
        queue_header.addWidget(self.queue_count_label)
        queue_header.addStretch()
        queue_header.addWidget(remove_btn)
        queue_header.addWidget(clear_btn)
        queue_header.addWidget(self.print_btn)
        queue_layout.addLayout(queue_header)

        self.queue_list = QListWidget()
        self.queue_list.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        self.queue_list.setMaximumHeight(160)
        queue_layout.addWidget(self.queue_list)

        splitter.addWidget(queue_widget)
        splitter.setSizes([420, 200])

        # status bar
        self.status = QStatusBar()
        self.setStatusBar(self.status)
        self._refresh_status()

    # ------------------------------------------------------------------
    # Slots
    # ------------------------------------------------------------------

    def _on_select_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select a folder", str(Path.home()))
        if not folder:
            return
        self._current_folder = Path(folder)
        self.folder_label.setText(folder)
        self.folder_label.setStyleSheet("")
        self._scan_folder()

    def _scan_folder(self):
        self.table.setSortingEnabled(False)
        self.table.setRowCount(0)

        if not self._current_folder:
            return

        images = [
            p for p in self._current_folder.rglob("*")
            if p.is_file() and p.suffix.lower() in IMAGE_EXTENSIONS
        ]

        for img in images:
            row = self.table.rowCount()
            self.table.insertRow(row)

            name_item = QTableWidgetItem(img.name)
            name_item.setToolTip(str(img))
            name_item.setData(Qt.ItemDataRole.UserRole, str(img))

            ext_item = QTableWidgetItem(img.suffix.lower())

            size_bytes = img.stat().st_size
            size_item = QTableWidgetItem(_human_size(size_bytes))
            size_item.setData(Qt.ItemDataRole.UserRole, size_bytes)

            self.table.setItem(row, 0, name_item)
            self.table.setItem(row, 1, ext_item)
            self.table.setItem(row, 2, size_item)

        self.table.setSortingEnabled(True)
        self._refresh_status()

    def _on_search(self, text: str):
        text = text.lower()
        for row in range(self.table.rowCount()):
            item = self.table.item(row, 0)
            visible = text in (item.text().lower() if item else "")
            self.table.setRowHidden(row, not visible if text else False)

    def _on_add_to_queue(self):
        selected_rows = self.table.selectionModel().selectedRows()
        if not selected_rows:
            self.status.showMessage("No images selected — click rows to select them first.", 3000)
            return

        added = 0
        for index in selected_rows:
            name_item = self.table.item(index.row(), 0)
            if not name_item:
                continue
            path = Path(name_item.data(Qt.ItemDataRole.UserRole))
            if path not in self._queue:
                self._queue.append(path)
                list_item = QListWidgetItem(f"{path.name}  ({path.parent})")
                list_item.setData(Qt.ItemDataRole.UserRole, str(path))
                list_item.setToolTip(str(path))
                self.queue_list.addItem(list_item)
                added += 1

        self._refresh_queue_ui()
        self.status.showMessage(f"Added {added} image(s) to queue. Queue total: {len(self._queue)}.", 3000)

    def _on_remove_from_queue(self):
        selected = self.queue_list.selectedItems()
        for item in selected:
            path = Path(item.data(Qt.ItemDataRole.UserRole))
            self._queue = [p for p in self._queue if p != path]
            self.queue_list.takeItem(self.queue_list.row(item))
        self._refresh_queue_ui()

    def _on_clear_queue(self):
        self._queue.clear()
        self.queue_list.clear()
        self._refresh_queue_ui()

    def _on_print_docx(self):
        if not self._queue:
            return

        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save DOCX grid",
            str(Path.cwd() / "image_grid.docx"),
            "Word Document (*.docx)",
        )
        if not save_path:
            return

        try:
            _build_docx_grid(self._queue, Path(save_path))
            QMessageBox.information(
                self,
                "Done",
                f"Saved {len(self._queue)} image(s) to:\n{save_path}",
            )
        except Exception as exc:
            QMessageBox.critical(self, "Export failed", str(exc))

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _refresh_status(self):
        folder_str = str(self._current_folder) if self._current_folder else "none"
        self.status.showMessage(
            f"{self.table.rowCount()} image(s) found  |  folder: {folder_str}"
        )

    def _refresh_queue_ui(self):
        count = len(self._queue)
        self.queue_count_label.setText(f"{count} image(s)")
        self.print_btn.setEnabled(count > 0)


# ------------------------------------------------------------------
# DOCX export
# ------------------------------------------------------------------

def _build_docx_grid(paths: list[Path], dest: Path) -> None:
    doc = Document()

    # Narrow margins to maximise usable page width
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    heading = doc.add_paragraph("Image Grid")
    heading.style = doc.styles["Heading 1"]

    # Build table
    cols = GRID_COLUMNS
    rows_needed = (len(paths) + cols - 1) // cols
    table = doc.add_table(rows=rows_needed, cols=cols)
    table.style = "Table Grid"

    for col_idx in range(cols):
        for cell in table.columns[col_idx].cells:
            cell.width = Inches(COL_WIDTH_IN)

    for i, img_path in enumerate(paths):
        row_idx = i // cols
        col_idx = i % cols
        cell = table.cell(row_idx, col_idx)
        cell.width = Inches(COL_WIDTH_IN)

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()

        img_bytes = _resize_image(img_path, THUMB_PX)
        run.add_picture(img_bytes, width=Inches(COL_WIDTH_IN - 0.1))

        # Filename caption below the image
        cap_para = cell.add_paragraph(img_path.name)
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.runs[0].font.size = Pt(7)

    # Clear leftover empty cells in the last row
    total_slots = rows_needed * cols
    for empty_slot in range(len(paths), total_slots):
        row_idx = empty_slot // cols
        col_idx = empty_slot % cols
        cell = table.cell(row_idx, col_idx)
        cell.paragraphs[0].clear()

    doc.save(str(dest))


def _resize_image(path: Path, max_px: int) -> io.BytesIO:
    """Open an image, thumbnail it to max_px×max_px, return as JPEG BytesIO."""
    with PilImage.open(path) as img:
        img = img.convert("RGB")
        img.thumbnail((max_px, max_px), PilImage.Resampling.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        buf.seek(0)
        return buf


# ------------------------------------------------------------------
# Misc helpers
# ------------------------------------------------------------------

def _human_size(n: int) -> str:
    for unit in ("B", "KB", "MB", "GB"):
        if n < 1024:
            return f"{n:.1f} {unit}"
        n /= 1024
    return f"{n:.1f} TB"
